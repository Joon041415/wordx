#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PyDocx Editor
포괄적인 기능을 갖춘 Python 기반 Word(.docx) 데스크톱 편집기.

필요 패키지:
    pip install python-docx pillow

실행:
    python app.py
"""

import os
import io
import uuid
import tempfile
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser

try:
    from PIL import Image, ImageTk
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# --------------------------------------------------------------------------
# 상수
# --------------------------------------------------------------------------

STYLE_NAMES = ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]
STYLE_PREVIEW_FONT = {
    "Normal":     ("맑은 고딕", 11, "normal"),
    "Title":      ("맑은 고딕", 28, "bold"),
    "Heading 1":  ("맑은 고딕", 20, "bold"),
    "Heading 2":  ("맑은 고딕", 16, "bold"),
    "Heading 3":  ("맑은 고딕", 13, "bold"),
    "Heading 4":  ("맑은 고딕", 12, "bold"),
}

FONT_FAMILIES = ["맑은 고딕", "Malgun Gothic", "Calibri", "Arial",
                  "Times New Roman", "Georgia", "Verdana", "Courier New"]
FONT_SIZES = [8, 9, 10, 10.5, 11, 12, 14, 16, 18, 20, 24, 28, 32, 36, 40, 48, 54, 60, 66, 72]

ALIGN_NAMES = ["left", "center", "right", "justify"]
ALIGN_TO_DOCX = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
DOCX_ALIGN_TO_NAME = {v: k for k, v in ALIGN_TO_DOCX.items()}

TABLE_MARK = "\u2b1f TABLE:{tid} \u2b1f"   # placeholder line text for embedded tables


def rgb_to_hex(rgb):
    if rgb is None:
        return None
    return "#%02x%02x%02x" % (rgb[0], rgb[1], rgb[2])


def hex_to_rgbcolor(hexstr):
    hexstr = hexstr.lstrip("#")
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))


# --------------------------------------------------------------------------
# 메인 애플리케이션
# --------------------------------------------------------------------------

class DocxEditorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PyDocx Editor - 새 문서")
        self.root.geometry("1000x750")

        self.filepath = None
        self.dirty = False
        self.photo_refs = {}      # image_id -> PhotoImage (thumbnail shown in editor)
        self.images = {}          # image_id -> {"path": full-res source path}
        self.tables = {}          # table_id -> list[list[str]]  (row-major cell text)
        self._tmpdir = tempfile.mkdtemp(prefix="pydocx_")

        self._build_menu()
        self._build_toolbar()
        self._build_text_area()
        self._build_statusbar()
        self._configure_base_tags()

        self._new_document(ask=False)
        self.text.bind("<<Modified>>", self._on_modified)
        self.text.bind("<KeyRelease>", self._update_status)
        self.text.bind("<ButtonRelease>", self._update_status)

        self.root.protocol("WM_DELETE_WINDOW", self._on_close)

    # ------------------------------------------------------------------
    # UI 구성
    # ------------------------------------------------------------------

    def _build_menu(self):
        menubar = tk.Menu(self.root)

        filemenu = tk.Menu(menubar, tearoff=0)
        filemenu.add_command(label="새 문서", accelerator="Ctrl+N", command=self._new_document)
        filemenu.add_command(label="열기...", accelerator="Ctrl+O", command=self._open_document)
        filemenu.add_command(label="저장", accelerator="Ctrl+S", command=self._save_document)
        filemenu.add_command(label="다른 이름으로 저장...", command=self._save_document_as)
        filemenu.add_separator()
        filemenu.add_command(label="종료", command=self._on_close)
        menubar.add_cascade(label="파일", menu=filemenu)

        editmenu = tk.Menu(menubar, tearoff=0)
        editmenu.add_command(label="실행 취소", accelerator="Ctrl+Z", command=lambda: self.text.edit_undo())
        editmenu.add_command(label="다시 실행", accelerator="Ctrl+Y", command=lambda: self.text.edit_redo())
        editmenu.add_separator()
        editmenu.add_command(label="찾기/바꾸기...", accelerator="Ctrl+F", command=self._open_find_replace)
        menubar.add_cascade(label="편집", menu=editmenu)

        insertmenu = tk.Menu(menubar, tearoff=0)
        insertmenu.add_command(label="이미지 삽입...", command=self._insert_image)
        insertmenu.add_command(label="표 삽입...", command=self._insert_table)
        menubar.add_cascade(label="삽입", menu=insertmenu)

        self.root.config(menu=menubar)
        self.root.bind_all("<Control-n>", lambda e: self._new_document())
        self.root.bind_all("<Control-o>", lambda e: self._open_document())
        self.root.bind_all("<Control-s>", lambda e: self._save_document())
        self.root.bind_all("<Control-f>", lambda e: self._open_find_replace())
        self.root.bind_all("<Control-b>", lambda e: self._toggle_tag("bold"))
        self.root.bind_all("<Control-i>", lambda e: self._toggle_tag("italic"))
        self.root.bind_all("<Control-u>", lambda e: self._toggle_tag("underline"))

    def _build_toolbar(self):
        bar = ttk.Frame(self.root)
        bar.pack(side="top", fill="x", padx=4, pady=4)

        # 문단 스타일
        ttk.Label(bar, text="스타일").pack(side="left", padx=(0, 2))
        self.style_var = tk.StringVar(value="Normal")
        style_box = ttk.Combobox(bar, textvariable=self.style_var, values=STYLE_NAMES,
                                  width=12, state="readonly")
        style_box.pack(side="left", padx=(0, 8))
        style_box.bind("<<ComboboxSelected>>", lambda e: self._apply_paragraph_style(self.style_var.get()))

        # 글꼴
        ttk.Label(bar, text="글꼴").pack(side="left", padx=(0, 2))
        self.font_var = tk.StringVar(value=FONT_FAMILIES[0])
        font_box = ttk.Combobox(bar, textvariable=self.font_var, values=FONT_FAMILIES,
                                 width=14, state="readonly")
        font_box.pack(side="left", padx=(0, 8))
        font_box.bind("<<ComboboxSelected>>", lambda e: self._apply_font_family(self.font_var.get()))

        # 크기
        ttk.Label(bar, text="크기").pack(side="left", padx=(0, 2))
        self.size_var = tk.StringVar(value="11")
        size_box = ttk.Combobox(bar, textvariable=self.size_var,
                                 values=[str(s) for s in FONT_SIZES], width=5, state="readonly")
        size_box.pack(side="left", padx=(0, 8))
        size_box.bind("<<ComboboxSelected>>", lambda e: self._apply_font_size(self.size_var.get()))

        # B I U
        ttk.Button(bar, text="B", width=3, command=lambda: self._toggle_tag("bold")).pack(side="left")
        ttk.Button(bar, text="I", width=3, command=lambda: self._toggle_tag("italic")).pack(side="left")
        ttk.Button(bar, text="U", width=3, command=lambda: self._toggle_tag("underline")).pack(side="left", padx=(0, 8))

        # 색상
        ttk.Button(bar, text="글자색", command=self._choose_color).pack(side="left", padx=(0, 8))

        # 정렬
        for name, label in [("left", "왼쪽"), ("center", "가운데"), ("right", "오른쪽"), ("justify", "양쪽")]:
            ttk.Button(bar, text=label, width=5,
                       command=lambda n=name: self._apply_alignment(n)).pack(side="left")

        ttk.Separator(bar, orient="vertical").pack(side="left", fill="y", padx=8)
        ttk.Button(bar, text="이미지 삽입", command=self._insert_image).pack(side="left", padx=2)
        ttk.Button(bar, text="표 삽입", command=self._insert_table).pack(side="left", padx=2)

    def _build_text_area(self):
        frame = ttk.Frame(self.root)
        frame.pack(fill="both", expand=True, padx=4, pady=(0, 4))

        yscroll = ttk.Scrollbar(frame, orient="vertical")
        self.text = tk.Text(frame, wrap="word", undo=True, font=("맑은 고딕", 11),
                             yscrollcommand=yscroll.set, padx=12, pady=12)
        yscroll.config(command=self.text.yview)
        yscroll.pack(side="right", fill="y")
        self.text.pack(side="left", fill="both", expand=True)

        # 표 placeholder를 더블클릭하면 편집기가 열리도록
        self.text.bind("<Double-Button-1>", self._maybe_edit_table_at_click)

    def _build_statusbar(self):
        self.status_var = tk.StringVar(value="준비됨")
        bar = ttk.Label(self.root, textvariable=self.status_var, anchor="w", relief="sunken")
        bar.pack(side="bottom", fill="x")

    def _configure_base_tags(self):
        self.text.tag_configure("bold", font=("맑은 고딕", 11, "bold"))
        self.text.tag_configure("italic", font=("맑은 고딕", 11, "italic"))
        self.text.tag_configure("underline", underline=True)
        for name, (fam, sz, weight) in STYLE_PREVIEW_FONT.items():
            self.text.tag_configure(f"style_{name}", font=(fam, sz, weight))
        for name in ALIGN_NAMES:
            tk_justify = "left" if name == "justify" else name
            self.text.tag_configure(f"align_{name}", justify=tk_justify)
        # 우선순위: 굵게/기울임/밑줄/글꼴/크기/색상이 문단 스타일보다 위에 오도록
        for name in STYLE_PREVIEW_FONT:
            self.text.tag_lower(f"style_{name}")

    # ------------------------------------------------------------------
    # 서식 적용 (선택 영역에 태그 적용)
    # ------------------------------------------------------------------

    def _selection_range(self):
        try:
            return self.text.index("sel.first"), self.text.index("sel.last")
        except tk.TclError:
            return None, None

    def _toggle_tag(self, tag):
        start, end = self._selection_range()
        if not start:
            return
        ranges = self.text.tag_ranges(tag)
        currently_on = self.text.tag_nextrange(tag, start, end)
        if currently_on:
            self.text.tag_remove(tag, start, end)
        else:
            self.text.tag_add(tag, start, end)
        self._mark_dirty()

    def _apply_exclusive_tag(self, prefix, new_tag, start, end, configure_fn=None):
        """prefix로 시작하는 기존 태그를 제거하고 new_tag만 남긴다 (같은 종류는 하나만 적용)."""
        for tag in list(self.text.tag_names()):
            if tag.startswith(prefix):
                self.text.tag_remove(tag, start, end)
        if configure_fn:
            configure_fn()
        self.text.tag_add(new_tag, start, end)

    def _apply_font_family(self, family):
        start, end = self._selection_range()
        if not start:
            return
        tag = f"font_{family}"

        def cfg():
            self.text.tag_configure(tag, font=(family, int(float(self.size_var.get()))))
        self._apply_exclusive_tag("font_", tag, start, end, cfg)
        self._mark_dirty()

    def _apply_font_size(self, size):
        start, end = self._selection_range()
        if not start:
            return
        tag = f"size_{size}"

        def cfg():
            self.text.tag_configure(tag, font=(self.font_var.get(), int(float(size))))
        self._apply_exclusive_tag("size_", tag, start, end, cfg)
        self._mark_dirty()

    def _choose_color(self):
        start, end = self._selection_range()
        if not start:
            return
        rgb, hexstr = colorchooser.askcolor(title="글자 색상 선택")
        if not hexstr:
            return
        tag = f"color_{hexstr.lstrip('#')}"

        def cfg():
            self.text.tag_configure(tag, foreground=hexstr)
        self._apply_exclusive_tag("color_", tag, start, end, cfg)
        self._mark_dirty()

    def _apply_alignment(self, name):
        line = self.text.index("insert").split(".")[0]
        start, end = f"{line}.0", f"{line}.end +1c"
        for a in ALIGN_NAMES:
            self.text.tag_remove(f"align_{a}", start, end)
        self.text.tag_add(f"align_{name}", start, end)
        self._mark_dirty()

    def _apply_paragraph_style(self, style_name):
        line = self.text.index("insert").split(".")[0]
        start, end = f"{line}.0", f"{line}.end +1c"
        for s in STYLE_PREVIEW_FONT:
            self.text.tag_remove(f"style_{s}", start, end)
        self.text.tag_add(f"style_{style_name}", start, end)
        self._mark_dirty()

    # ------------------------------------------------------------------
    # 이미지 삽입
    # ------------------------------------------------------------------

    def _insert_image(self):
        path = filedialog.askopenfilename(
            title="삽입할 이미지 선택",
            filetypes=[("이미지 파일", "*.png *.jpg *.jpeg *.gif *.bmp"), ("모든 파일", "*.*")])
        if not path:
            return
        image_id = uuid.uuid4().hex
        thumb = self._make_thumbnail(path)
        if thumb is None:
            messagebox.showerror("오류", "이미지를 불러올 수 없습니다 (Pillow가 필요할 수 있습니다).")
            return
        self.photo_refs[image_id] = thumb
        self.images[image_id] = {"path": path}
        self.text.image_create(tk.INSERT, image=thumb, name=image_id)
        self._mark_dirty()

    def _make_thumbnail(self, path, max_w=400):
        try:
            if PIL_AVAILABLE:
                im = Image.open(path)
                im.load()
                if im.mode not in ("RGB", "RGBA"):
                    im = im.convert("RGBA")
                w, h = im.size
                if w > max_w:
                    ratio = max_w / float(w)
                    im = im.resize((max_w, int(h * ratio)))
                return ImageTk.PhotoImage(im)
            else:
                return tk.PhotoImage(file=path)
        except Exception:
            return None

    # ------------------------------------------------------------------
    # 표 삽입 / 편집
    # ------------------------------------------------------------------

    def _insert_table(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("표 삽입")
        dlg.transient(self.root)
        ttk.Label(dlg, text="행 수:").grid(row=0, column=0, padx=6, pady=6)
        rows_var = tk.IntVar(value=3)
        ttk.Spinbox(dlg, from_=1, to=50, textvariable=rows_var, width=6).grid(row=0, column=1)
        ttk.Label(dlg, text="열 수:").grid(row=1, column=0, padx=6, pady=6)
        cols_var = tk.IntVar(value=3)
        ttk.Spinbox(dlg, from_=1, to=20, textvariable=cols_var, width=6).grid(row=1, column=1)

        def create():
            rows, cols = rows_var.get(), cols_var.get()
            data = [["" for _ in range(cols)] for _ in range(rows)]
            table_id = uuid.uuid4().hex
            self.tables[table_id] = data
            dlg.destroy()
            self._insert_table_placeholder(table_id)
            self._open_table_editor(table_id)

        ttk.Button(dlg, text="만들기", command=create).grid(row=2, column=0, columnspan=2, pady=8)

    def _insert_table_placeholder(self, table_id):
        line = self.text.index("insert").split(".")[0]
        if self.text.get(f"{line}.0", f"{line}.end").strip() != "":
            self.text.insert("insert", "\n")
        mark = TABLE_MARK.format(tid=table_id)
        self.text.insert("insert", mark + "\n")
        newline = self.text.index("insert").split(".")[0]
        placeholder_line = str(int(newline) - 1)
        self.text.tag_add("table_placeholder", f"{placeholder_line}.0", f"{placeholder_line}.end")
        self.text.tag_configure("table_placeholder", background="#eef3ff", foreground="#3355aa")
        self._mark_dirty()

    def _maybe_edit_table_at_click(self, event):
        index = self.text.index(f"@{event.x},{event.y}")
        line = index.split(".")[0]
        line_text = self.text.get(f"{line}.0", f"{line}.end")
        if line_text.strip().startswith("\u2b1f TABLE:"):
            table_id = line_text.strip().split("TABLE:")[1].rsplit("\u2b1f", 1)[0].strip()
            if table_id in self.tables:
                self._open_table_editor(table_id)

    def _open_table_editor(self, table_id):
        data = self.tables[table_id]
        dlg = tk.Toplevel(self.root)
        dlg.title("표 편집")
        entries = []
        for r, row in enumerate(data):
            row_entries = []
            for c, val in enumerate(row):
                e = ttk.Entry(dlg, width=14)
                e.insert(0, val)
                e.grid(row=r, column=c, padx=1, pady=1)
                row_entries.append(e)
            entries.append(row_entries)

        def save_and_close():
            for r, row_entries in enumerate(entries):
                for c, e in enumerate(row_entries):
                    data[r][c] = e.get()
            self._mark_dirty()
            dlg.destroy()

        ttk.Button(dlg, text="저장", command=save_and_close).grid(
            row=len(data), column=0, columnspan=max(1, len(data[0]) if data else 1), pady=6)

    # ------------------------------------------------------------------
    # 찾기 / 바꾸기
    # ------------------------------------------------------------------

    def _open_find_replace(self):
        dlg = tk.Toplevel(self.root)
        dlg.title("찾기/바꾸기")
        ttk.Label(dlg, text="찾을 내용:").grid(row=0, column=0, padx=6, pady=6, sticky="e")
        find_var = tk.StringVar()
        ttk.Entry(dlg, textvariable=find_var, width=30).grid(row=0, column=1, padx=6, pady=6)
        ttk.Label(dlg, text="바꿀 내용:").grid(row=1, column=0, padx=6, pady=6, sticky="e")
        replace_var = tk.StringVar()
        ttk.Entry(dlg, textvariable=replace_var, width=30).grid(row=1, column=1, padx=6, pady=6)

        def do_find():
            self.text.tag_remove("search_hl", "1.0", "end")
            term = find_var.get()
            if not term:
                return
            start = "1.0"
            count = 0
            while True:
                pos = self.text.search(term, start, stopindex="end")
                if not pos:
                    break
                end = f"{pos}+{len(term)}c"
                self.text.tag_add("search_hl", pos, end)
                start = end
                count += 1
            self.text.tag_configure("search_hl", background="yellow")
            self.status_var.set(f"{count}개 찾음")

        def do_replace_all():
            term, repl = find_var.get(), replace_var.get()
            if not term:
                return
            start = "1.0"
            count = 0
            while True:
                pos = self.text.search(term, start, stopindex="end")
                if not pos:
                    break
                end = f"{pos}+{len(term)}c"
                self.text.delete(pos, end)
                self.text.insert(pos, repl)
                start = f"{pos}+{len(repl)}c"
                count += 1
            self.status_var.set(f"{count}개 바꿈")
            self._mark_dirty()

        ttk.Button(dlg, text="모두 찾기", command=do_find).grid(row=2, column=0, pady=6)
        ttk.Button(dlg, text="모두 바꾸기", command=do_replace_all).grid(row=2, column=1, pady=6)

    # ------------------------------------------------------------------
    # 문서 생성 / 열기 / 저장
    # ------------------------------------------------------------------

    def _new_document(self, ask=True):
        if ask and not self._confirm_discard_changes():
            return
        self.text.delete("1.0", "end")
        self.images.clear()
        self.tables.clear()
        self.photo_refs.clear()
        self.filepath = None
        self.dirty = False
        self.text.tag_add("style_Normal", "1.0", "end")
        self.root.title("PyDocx Editor - 새 문서")
        self._update_status()

    def _confirm_discard_changes(self):
        if not self.dirty:
            return True
        res = messagebox.askyesnocancel("저장하지 않은 변경사항", "변경사항을 저장하시겠습니까?")
        if res is None:
            return False
        if res:
            self._save_document()
        return True

    def _open_document(self):
        if not self._confirm_discard_changes():
            return
        path = filedialog.askopenfilename(
            title="Word 문서 열기", filetypes=[("Word 문서", "*.docx")])
        if not path:
            return
        try:
            self._load_docx(path)
        except Exception as exc:
            messagebox.showerror("열기 실패", f"문서를 여는 중 오류가 발생했습니다:\n{exc}")
            return
        self.filepath = path
        self.dirty = False
        self.root.title(f"PyDocx Editor - {os.path.basename(path)}")
        self._update_status()

    def _load_docx(self, path):
        doc = Document(path)
        self.text.delete("1.0", "end")
        self.images.clear()
        self.tables.clear()
        self.photo_refs.clear()

        body = doc.element.body
        for child in body.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):
                self._load_paragraph(doc, child)
            elif tag == qn("w:tbl"):
                self._load_table(doc, child)
        # trailing sectPr etc. ignored

        if self.text.get("1.0", "end").strip() == "" and not self.tables:
            self.text.insert("1.0", "\n")

    def _load_paragraph(self, doc, p_element):
        from docx.text.paragraph import Paragraph
        para = Paragraph(p_element, doc)
        line_start = self.text.index("end-1c")

        style_name = para.style.name if para.style else "Normal"
        if style_name not in STYLE_PREVIEW_FONT:
            style_name = "Normal"

        for run in para.runs:
            self._insert_run(doc, run)

        self.text.insert("end", "\n")
        line_no = line_start.split(".")[0]
        line_end = self.text.index("end-1c").split(".")[0]
        self.text.tag_add(f"style_{style_name}", f"{line_no}.0", f"{line_end}.0")

        align_name = DOCX_ALIGN_TO_NAME.get(para.alignment)
        if align_name:
            self.text.tag_add(f"align_{align_name}", f"{line_no}.0", f"{line_end}.0")

    def _insert_run(self, doc, run):
        # 이미지가 포함된 run 처리
        drawings = run._element.findall(".//" + qn("w:drawing"))
        if drawings:
            for drawing in drawings:
                blip = drawing.find(".//" + qn("a:blip"))
                if blip is not None:
                    rId = blip.get(qn("r:embed"))
                    try:
                        part = doc.part.related_parts[rId]
                        tmp_path = os.path.join(self._tmpdir, uuid.uuid4().hex + ".png")
                        with open(tmp_path, "wb") as f:
                            f.write(part.blob)
                        image_id = uuid.uuid4().hex
                        thumb = self._make_thumbnail(tmp_path)
                        if thumb is not None:
                            self.photo_refs[image_id] = thumb
                            self.images[image_id] = {"path": tmp_path}
                            self.text.image_create("end", image=thumb, name=image_id)
                    except Exception:
                        pass
            return

        text = run.text
        if not text:
            return
        start = self.text.index("end-1c")
        self.text.insert("end", text)
        end = self.text.index("end-1c")

        if run.bold:
            self.text.tag_add("bold", start, end)
        if run.italic:
            self.text.tag_add("italic", start, end)
        if run.underline:
            self.text.tag_add("underline", start, end)
        if run.font.name:
            fam = run.font.name
            tag = f"font_{fam}"
            self.text.tag_configure(tag, font=(fam, int(run.font.size.pt) if run.font.size else 11))
            self.text.tag_add(tag, start, end)
        if run.font.size:
            sz = run.font.size.pt
            tag = f"size_{sz}"
            fam = run.font.name or self.font_var.get()
            self.text.tag_configure(tag, font=(fam, int(sz)))
            self.text.tag_add(tag, start, end)
        if run.font.color and run.font.color.rgb:
            hexstr = "#" + str(run.font.color.rgb)
            tag = f"color_{str(run.font.color.rgb)}"
            self.text.tag_configure(tag, foreground=hexstr)
            self.text.tag_add(tag, start, end)

    def _load_table(self, doc, tbl_element):
        from docx.table import Table
        table = Table(tbl_element, doc)
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        table_id = uuid.uuid4().hex
        self.tables[table_id] = data
        mark = TABLE_MARK.format(tid=table_id)
        line_start = self.text.index("end-1c").split(".")[0]
        self.text.insert("end", mark + "\n")
        self.text.tag_add("table_placeholder", f"{line_start}.0", f"{line_start}.end")
        self.text.tag_configure("table_placeholder", background="#eef3ff", foreground="#3355aa")

    def _save_document(self):
        if self.filepath:
            self._write_docx(self.filepath)
        else:
            self._save_document_as()

    def _save_document_as(self):
        path = filedialog.asksaveasfilename(
            title="다른 이름으로 저장", defaultextension=".docx",
            filetypes=[("Word 문서", "*.docx")])
        if not path:
            return
        self._write_docx(path)
        self.filepath = path
        self.root.title(f"PyDocx Editor - {os.path.basename(path)}")

    def _write_docx(self, path):
        try:
            doc = Document()
            last_line = int(self.text.index("end-1c").split(".")[0])

            for line_no in range(1, last_line + 1):
                line_text = self.text.get(f"{line_no}.0", f"{line_no}.end")

                if line_text.strip().startswith("\u2b1f TABLE:"):
                    table_id = line_text.strip().split("TABLE:")[1].rsplit("\u2b1f", 1)[0].strip()
                    if table_id in self.tables:
                        self._write_table(doc, self.tables[table_id])
                    continue

                self._write_paragraph(doc, line_no)

            doc.save(path)
            self.dirty = False
            self.status_var.set(f"저장됨: {path}")
        except Exception as exc:
            messagebox.showerror("저장 실패", f"문서를 저장하는 중 오류가 발생했습니다:\n{exc}")

    def _write_paragraph(self, doc, line_no):
        style_name = "Normal"
        for s in STYLE_PREVIEW_FONT:
            if self.text.tag_nextrange(f"style_{s}", f"{line_no}.0", f"{line_no}.end"):
                style_name = s
                break

        align_name = None
        for a in ALIGN_NAMES:
            if self.text.tag_nextrange(f"align_{a}", f"{line_no}.0", f"{line_no}.end"):
                align_name = a
                break

        para = doc.add_paragraph(style=style_name if style_name != "Normal" else None)
        if align_name:
            para.alignment = ALIGN_TO_DOCX[align_name]

        for seg in self._extract_line_segments(line_no):
            if seg["type"] == "image":
                run = para.add_run()
                img_info = self.images.get(seg["id"])
                if img_info:
                    try:
                        run.add_picture(img_info["path"], width=Emu(4000000))
                    except Exception:
                        pass
            else:
                if seg["text"] == "":
                    continue
                run = para.add_run(seg["text"])
                tags = seg["tags"]
                if "bold" in tags:
                    run.bold = True
                if "italic" in tags:
                    run.italic = True
                if "underline" in tags:
                    run.underline = True
                for t in tags:
                    if t.startswith("font_"):
                        run.font.name = t[len("font_"):]
                    elif t.startswith("size_"):
                        try:
                            run.font.size = Pt(float(t[len("size_"):]))
                        except ValueError:
                            pass
                    elif t.startswith("color_"):
                        try:
                            run.font.color.rgb = hex_to_rgbcolor(t[len("color_"):])
                        except Exception:
                            pass

    def _write_table(self, doc, data):
        if not data or not data[0]:
            return
        rows, cols = len(data), len(data[0])
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        for r in range(rows):
            for c in range(cols):
                table.cell(r, c).text = data[r][c]

    def _extract_line_segments(self, line_no):
        start, end = f"{line_no}.0", f"{line_no}.end"
        dump = self.text.dump(start, end, text=True, tag=True, image=True)
        active = set()
        segments = []
        FORMAT_PREFIXES = ("bold", "italic", "underline", "font_", "size_", "color_")

        def is_format_tag(name):
            return name in ("bold", "italic", "underline") or name.startswith(("font_", "size_", "color_"))

        for key, value, index in dump:
            if key == "tagon" and is_format_tag(value):
                active.add(value)
            elif key == "tagoff" and value in active:
                active.discard(value)
            elif key == "text":
                segments.append({"type": "text", "text": value, "tags": set(active)})
            elif key == "image":
                segments.append({"type": "image", "id": value})
        return segments

    # ------------------------------------------------------------------
    # 상태 관리
    # ------------------------------------------------------------------

    def _on_modified(self, event=None):
        if self.text.edit_modified():
            self.dirty = True
            self.text.edit_modified(False)

    def _mark_dirty(self):
        self.dirty = True

    def _update_status(self, event=None):
        content = self.text.get("1.0", "end-1c")
        words = len(content.split())
        chars = len(content)
        line, col = self.text.index("insert").split(".")
        name = os.path.basename(self.filepath) if self.filepath else "새 문서"
        mark = " *" if self.dirty else ""
        self.status_var.set(f"{name}{mark}   |   단어 수: {words}   |   글자 수: {chars}   |   {line}행 {int(col)+1}열")

    def _on_close(self):
        if self._confirm_discard_changes():
            self.root.destroy()


def main():
    root = tk.Tk()
    app = DocxEditorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
